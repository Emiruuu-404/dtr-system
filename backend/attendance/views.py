import json
import re
from django.http import JsonResponse, FileResponse
from django.views.decorators.csrf import csrf_exempt
from django.utils import timezone
from .models import Intern, Attendance, AccomplishmentReport, AccomplishmentImage
from django.contrib.auth.hashers import make_password, check_password
import io
from datetime import datetime, time, timedelta
import math

def get_effective_hours(start_dt, end_dt):
    if not start_dt or not end_dt:
        return 0
    sec = (end_dt - start_dt).total_seconds()
    if sec <= 0:
        return 0

    lunch_start = timezone.make_aware(datetime.combine(start_dt.date(), time(12, 0)))
    lunch_end = timezone.make_aware(datetime.combine(start_dt.date(), time(13, 0)))

    overlap_start = max(start_dt, lunch_start)
    overlap_end = min(end_dt, lunch_end)
    
    if overlap_start < overlap_end:
        sec -= (overlap_end - overlap_start).total_seconds()

    return sec / 3600

def format_hrs_mins(decimal_hours):
    if decimal_hours <= 0:
        return "0 HRS"
    hrs = int(decimal_hours)
    mins = int(round((decimal_hours - hrs) * 60))
    if hrs > 0 and mins > 0:
        return f"{hrs} hrs and {mins} mins"
    elif hrs > 0:
        return f"{hrs} hrs"
    else:
        return f"{mins} mins"

@csrf_exempt
def register(request):

    if request.method != "POST":
        return JsonResponse({"error": "POST request required"}, status=405)

    try:
        data = json.loads(request.body)
    except:
        return JsonResponse({"error": "Invalid JSON"}, status=400)

    name = data.get("name")
    student_id = data.get("student_id")
    email = data.get("email")
    password = data.get("password")

    if not all([name, student_id, email, password]):
        return JsonResponse({"error": "Missing fields"}, status=400)

    # validate student id format XX-XXXX
    if not re.match(r"^\d{2}-\d{4}$", student_id):
        return JsonResponse({"error": "Student ID format must be XX-XXXX"}, status=400)

    if Intern.objects.filter(student_id=student_id).exists():
        return JsonResponse({"error": "Student ID already registered"}, status=400)

    user = Intern.objects.create(
        name=name,
        student_id=student_id,
        email=email,
        password=make_password(password)
    )
    
    return JsonResponse({"message": "Account created successfully"})


@csrf_exempt
def login_view(request):
    if request.method != "POST":
        return JsonResponse({"error": "POST request required"}, status=405)

    try:
        data = json.loads(request.body)
    except:
        return JsonResponse({"error": "Invalid JSON"}, status=400)

    student_id = data.get("student_id")
    password = data.get("password")

    if not all([student_id, password]):
        return JsonResponse({"error": "Missing fields"}, status=400)

    try:
        if "@" in student_id:
            user = Intern.objects.get(email=student_id)
        else:
            user = Intern.objects.get(student_id=student_id)
            
        print(f"DEBUG LOGIN: Received password: '{password}', stored hash: '{user.password}'")
        if check_password(password, user.password):
            return JsonResponse({
                "message": "Login successful",
                "student_id": user.student_id,
                "name": user.name
            })
        else:
            print(f"Unauthorized: Password mismatch for user {user.student_id}")
            return JsonResponse({"error": "Invalid credentials"}, status=401)
    except Intern.DoesNotExist:
        print(f"Unauthorized: User not found for login ID {student_id}")
        return JsonResponse({"error": "Invalid credentials"}, status=401)

@csrf_exempt
def time_in(request):

    if request.method != "POST":
        return JsonResponse({"error": "POST required"}, status=405)

    data = json.loads(request.body)
    student_id = data.get("student_id")

    now = timezone.localtime()
    today = now.date()
    current_hour = now.hour

    # Block if past 5 PM (17:00)
    if current_hour >= 17:
        return JsonResponse({
            "error": "OJT hours has already ended."
        }, status=400)

    # Check existing record today
    existing = Attendance.objects.filter(
        student_id=student_id,
        date=today
    ).first()

    # 🚫 Block if already timed in today
    if existing and existing.am_time_in:
        return JsonResponse({
            "error": "You have already timed in for today."
        }, status=400)

    # If no record today → create new
    if not existing:
        Attendance.objects.create(
            student_id=student_id,
            date=today,
            am_time_in=now
        )
    else:
        # If record exists but no time_in yet
        if not existing.am_time_in:
            existing.am_time_in = now
            existing.save()

    return JsonResponse({
        "message": "Time in recorded"
    })

@csrf_exempt
def time_out(request):
    if request.method != "POST":
        return JsonResponse({"error": "POST request required"}, status=405)

    try:
        data = json.loads(request.body)
    except:
        return JsonResponse({"error": "Invalid JSON"}, status=400)

    student_id = data.get("student_id")
    if not student_id:
        return JsonResponse({"error": "Missing student_id"}, status=400)

    try:
        intern = Intern.objects.get(student_id=student_id)
    except Intern.DoesNotExist:
        return JsonResponse({"error": "Intern not found"}, status=404)

    now = timezone.localtime()
    if now.hour >= 17:
        return JsonResponse({
            "error": "OJT hours has already ended."
        }, status=400)

    # In our new logic, time-out is not really used because we default hours,
    # but we still block it if needed.
    return JsonResponse({"error": "Time Out is automatically handled by the system today."}, status=400)




def get_leaderboards(request):
    interns = Intern.objects.all()
    
    leaderboard_data = []
    for intern in interns:
        records = Attendance.objects.filter(student_id=intern.student_id)
        total_hours = 0
        for r in records:
            total_hours += get_effective_hours(r.am_time_in, r.am_time_out)
            total_hours += get_effective_hours(r.pm_time_in, r.pm_time_out)
        total_hours = round(total_hours, 2)
        
        leaderboard_data.append({
            "id": intern.id,
            "name": intern.name,
            "hours": total_hours
        })
        
    # Sort by descending hours
    leaderboard_data.sort(key=lambda x: x['hours'], reverse=True)
    
    # Assign ranks
    for index, item in enumerate(leaderboard_data):
        item['rank'] = index + 1
        
    return JsonResponse({
        "leaderboard": leaderboard_data
    })

def get_status(request):
    student_id = request.GET.get("student_id")
    if not student_id:
        return JsonResponse({"error": "Missing student_id"}, status=400)

    try:
        user = Intern.objects.get(student_id=student_id)
    except Intern.DoesNotExist:
        return JsonResponse({"error": "Intern not found"}, status=404)

    today = timezone.localdate()

    record = Attendance.objects.filter(
        student_id=student_id,
        date=today
    ).first()

    # ===== STATUS =====
    if not record:
        status = "Not Timed In"
    elif record.am_time_in and not record.am_time_out:
        status = "AM IN"
    elif record.am_time_out and not record.pm_time_in:
        status = "AM OUT"
    elif record.pm_time_in and not record.pm_time_out:
        status = "PM IN"
    elif record.pm_time_out:
        status = "PM OUT"
    else:
        status = "Not Timed In"

    # ===== LAST TIME IN =====
    last_record = Attendance.objects.filter(
        student_id=student_id
    ).order_by('-date').first()

    if last_record and last_record.am_time_in:
        last_time = timezone.localtime(last_record.am_time_in)
        last_time_in = last_time.strftime("%I:%M %p")
    else:
        last_time_in = "--:--"

    # ===== TOTAL HOURS =====
    def compute_hours(r):
        total = 0
        total += get_effective_hours(r.am_time_in, r.am_time_out)
        total += get_effective_hours(r.pm_time_in, r.pm_time_out)
        return total

    all_records = Attendance.objects.filter(student_id=student_id)
    total_hours = round(sum(compute_hours(r) for r in all_records), 2)

    # ===== TODAY LOGS =====
    def fmt(t):
        return timezone.localtime(t).strftime("%I:%M %p") if t else "--:--"

    if record:
        today_logs = [
            {
                "in": fmt(record.am_time_in),
                "out": fmt(record.am_time_out),
                "in_label": "AM IN",
                "out_label": "AM OUT"
            },
            {
                "in": fmt(record.pm_time_in),
                "out": fmt(record.pm_time_out),
                "in_label": "PM IN",
                "out_label": "PM OUT"
            }
        ]
    else:
        today_logs = [
            {"in": "--:--", "out": "--:--", "in_label": "AM IN", "out_label": "AM OUT"},
            {"in": "--:--", "out": "--:--", "in_label": "PM IN", "out_label": "PM OUT"}
        ]

    # ===== ESTIMATED END DATE =====
    total_required = 486
    remaining_hours = max(total_required - total_hours, 0)
    
    if remaining_hours == 0:
        est_date_str = "Completed"
    else:
        remaining_days = math.ceil(remaining_hours / 8)
        est = today
        added_days = 0
        while added_days < remaining_days:
            est += timedelta(days=1)
            if est.weekday() < 5:
                added_days += 1
        est_date_str = est.strftime("%b %d, %Y")

    return JsonResponse({
        "name": user.name.split()[0],
        "status": status,
        "last_time_in": last_time_in,
        "today_logs": today_logs,
        "total_hours": total_hours,
        "formatted_total_hours": format_hrs_mins(total_hours),
        "total_required": total_required,
        "est_end_date": est_date_str
    })


def fmt(t):
    return timezone.localtime(t).strftime("%I:%M %p") if t else "--:--"

def get_history(request):
    student_id = request.GET.get("student_id")
    if not student_id:
        return JsonResponse({"error": "Missing student_id"}, status=400)

    records = Attendance.objects.filter(
        student_id=student_id
    ).order_by('-date')

    history = []

    for r in records:
        # Calculate hours
        total_hours = 0
        total_hours += get_effective_hours(r.am_time_in, r.am_time_out)
        total_hours += get_effective_hours(r.pm_time_in, r.pm_time_out)
        total_hours = round(total_hours, 2)

        # Determine first in and last out
        first_in = fmt(r.am_time_in) if r.am_time_in else fmt(r.pm_time_in)
        last_out = fmt(r.pm_time_out) if r.pm_time_out else fmt(r.am_time_out)

        # Status
        if r.am_time_in and r.am_time_out and r.pm_time_in and r.pm_time_out:
            status = "Completed"
        elif r.am_time_in or r.pm_time_in:
            status = "Incomplete"
        else:
            status = "No Data"

        history.append({
            "id": r.id,
            "date": r.date.strftime("%b %d, %Y"),
            "am_in": fmt(r.am_time_in),
            "am_out": fmt(r.am_time_out),
            "pm_in": fmt(r.pm_time_in),
            "pm_out": fmt(r.pm_time_out),
            "in": first_in,
            "out": last_out,
            "hours": total_hours,
            "formatted_hours": format_hrs_mins(total_hours),
            "status": status,
        })

    return JsonResponse({"records": history})

@csrf_exempt
def add_past_record(request):
    if request.method != "POST":
        return JsonResponse({"error": "POST required"}, status=405)

    try:
        data = json.loads(request.body)
    except:
        return JsonResponse({"error": "Invalid JSON"}, status=400)

    student_id = data.get("student_id")
    date_str = data.get("date")
    am_in_str = data.get("am_in")
    am_out_str = data.get("am_out")
    pm_in_str = data.get("pm_in")
    pm_out_str = data.get("pm_out")

    if not all([student_id, date_str]):
        return JsonResponse({"error": "Missing student_id or date"}, status=400)
    
    if not any([am_in_str, am_out_str, pm_in_str, pm_out_str]):
         return JsonResponse({"error": "Provide at least one time punch"}, status=400)

    try:
        intern = Intern.objects.get(student_id=student_id)
    except Intern.DoesNotExist:
        return JsonResponse({"error": "Intern not found"}, status=404)

    try:
        from datetime import datetime
        
        # Parse strings (expecting YYYY-MM-DD and HH:MM)
        date_obj = datetime.strptime(date_str, "%Y-%m-%d").date()
        
        am_in_obj = datetime.strptime(am_in_str, "%H:%M").time() if am_in_str else None
        am_out_obj = datetime.strptime(am_out_str, "%H:%M").time() if am_out_str else None
        pm_in_obj = datetime.strptime(pm_in_str, "%H:%M").time() if pm_in_str else None
        pm_out_obj = datetime.strptime(pm_out_str, "%H:%M").time() if pm_out_str else None

        # Determine time_in and time_out bounds for calculating hours and sorting
        aware_time_in = None
        if am_in_obj:
            aware_time_in = timezone.make_aware(datetime.combine(date_obj, am_in_obj))
        elif pm_in_obj:
            aware_time_in = timezone.make_aware(datetime.combine(date_obj, pm_in_obj))
            
        aware_time_out = None
        if pm_out_obj:
            aware_time_out = timezone.make_aware(datetime.combine(date_obj, pm_out_obj))
        elif am_out_obj:
            aware_time_out = timezone.make_aware(datetime.combine(date_obj, am_out_obj))

        record, _ = Attendance.objects.get_or_create(
            student_id=student_id,
            date=date_obj
        )
        
        if am_in_obj:
            record.am_time_in = timezone.make_aware(datetime.combine(date_obj, am_in_obj))
        if am_out_obj:
            record.am_time_out = timezone.make_aware(datetime.combine(date_obj, am_out_obj))
        if pm_in_obj:
            record.pm_time_in = timezone.make_aware(datetime.combine(date_obj, pm_in_obj))
        if pm_out_obj:
            record.pm_time_out = timezone.make_aware(datetime.combine(date_obj, pm_out_obj))
            
        record.save()

        return JsonResponse({"message": "Past record added successfully"})
    except Exception as e:
        return JsonResponse({"error": f"Error parsing datetime: {str(e)}"}, status=400)


@csrf_exempt
def save_today_record(request):
    if request.method != "POST":
        return JsonResponse({"error": "POST required"}, status=405)

    try:
        data = json.loads(request.body)
    except:
        return JsonResponse({"error": "Invalid JSON"}, status=400)

    student_id = data.get("student_id")
    if not student_id:
        return JsonResponse({"error": "Missing student_id"}, status=400)

    am_in_str = data.get("am_in")
    am_out_str = data.get("am_out")
    pm_in_str = data.get("pm_in")
    pm_out_str = data.get("pm_out")

    try:
        intern = Intern.objects.get(student_id=student_id)
    except Intern.DoesNotExist:
        return JsonResponse({"error": "Intern not found"}, status=404)

    try:
        from datetime import datetime
        
        today = timezone.localtime().date()

        # Helper to safely parse or set to None
        def parse_time(time_str):
            if not time_str or time_str == "--:--":
                return None
            if "AM" in time_str.upper() or "PM" in time_str.upper():
                 try:
                     return datetime.strptime(time_str, "%I:%M %p").time()
                 except ValueError:
                     pass
            try:
                return datetime.strptime(time_str, "%H:%M").time()
            except ValueError:
                return None

        am_in_obj = parse_time(am_in_str)
        am_out_obj = parse_time(am_out_str)
        pm_in_obj = parse_time(pm_in_str)
        pm_out_obj = parse_time(pm_out_str)

        aware_time_in = None
        if am_in_obj:
            aware_time_in = timezone.make_aware(datetime.combine(today, am_in_obj))
        elif pm_in_obj:
            aware_time_in = timezone.make_aware(datetime.combine(today, pm_in_obj))
            
        aware_time_out = None
        if pm_out_obj:
            aware_time_out = timezone.make_aware(datetime.combine(today, pm_out_obj))
        elif am_out_obj:
            aware_time_out = timezone.make_aware(datetime.combine(today, am_out_obj))

        # Check existing record today
        existing = Attendance.objects.filter(
            student_id=student_id,
            date=today
        ).first()

        if not existing:
            # Maybe there's a record for today based on am_time_in
            existing = Attendance.objects.filter(
                student_id=student_id,
                am_time_in__date=today
            ).first()

        if existing:
            existing.am_time_in = timezone.make_aware(datetime.combine(today, am_in_obj)) if am_in_obj else existing.am_time_in
            existing.am_time_out = timezone.make_aware(datetime.combine(today, am_out_obj)) if am_out_obj else existing.am_time_out
            existing.pm_time_in = timezone.make_aware(datetime.combine(today, pm_in_obj)) if pm_in_obj else existing.pm_time_in
            existing.pm_time_out = timezone.make_aware(datetime.combine(today, pm_out_obj)) if pm_out_obj else existing.pm_time_out
            existing.save()
        else:
            Attendance.objects.create(
                student_id=student_id,
                date=today,
                am_time_in=timezone.make_aware(datetime.combine(today, am_in_obj)) if am_in_obj else None,
                am_time_out=timezone.make_aware(datetime.combine(today, am_out_obj)) if am_out_obj else None,
                pm_time_in=timezone.make_aware(datetime.combine(today, pm_in_obj)) if pm_in_obj else None,
                pm_time_out=timezone.make_aware(datetime.combine(today, pm_out_obj)) if pm_out_obj else None,
            )

        return JsonResponse({"message": "Today's record updated successfully"})
    except Exception as e:
        return JsonResponse({"error": f"Error updating record: {str(e)}"}, status=400)




@csrf_exempt
def edit_record(request):
    if request.method != "POST":
        return JsonResponse({"error": "POST required"}, status=405)

    try:
        data = json.loads(request.body)
    except:
        return JsonResponse({"error": "Invalid JSON"}, status=400)

    student_id = data.get("student_id")
    record_id = data.get("record_id")
    
    if not all([student_id, record_id]):
        return JsonResponse({"error": "Missing student_id or record_id"}, status=400)

    am_in_str = data.get("am_in")
    am_out_str = data.get("am_out")
    pm_in_str = data.get("pm_in")
    pm_out_str = data.get("pm_out")

    try:
        record = Attendance.objects.get(id=record_id, student_id=student_id)
    except Attendance.DoesNotExist:
        return JsonResponse({"error": "Record not found"}, status=404)

    try:
        from datetime import datetime
        
        # Helper to safely parse or set to None
        def parse_time(time_str):
            if not time_str or time_str == "--:--":
                return None
            
            # Also handle time values containing AM/PM if they come from UI
            if "AM" in time_str.upper() or "PM" in time_str.upper():
                 try:
                     return datetime.strptime(time_str, "%I:%M %p").time()
                 except ValueError:
                     pass
            
            try:
                return datetime.strptime(time_str, "%H:%M").time()
            except ValueError:
                return None

        am_in_obj = parse_time(am_in_str)
        am_out_obj = parse_time(am_out_str)
        pm_in_obj = parse_time(pm_in_str)
        pm_out_obj = parse_time(pm_out_str)

        date_obj = record.date
        if not date_obj and record.am_time_in:
             date_obj = timezone.localtime(record.am_time_in).date()
             
        if not date_obj:
             date_obj = timezone.localtime().date() # Fallback

        # Determine time_in and time_out bounds for calculating hours
        aware_time_in = None
        if am_in_obj:
            aware_time_in = timezone.make_aware(datetime.combine(date_obj, am_in_obj))
        elif pm_in_obj:
            aware_time_in = timezone.make_aware(datetime.combine(date_obj, pm_in_obj))
            
        aware_time_out = None
        if pm_out_obj:
            aware_time_out = timezone.make_aware(datetime.combine(date_obj, pm_out_obj))
        elif am_out_obj:
            aware_time_out = timezone.make_aware(datetime.combine(date_obj, am_out_obj))

        # Get existing values
        old_am_in = record.am_time_in
        old_am_out = record.am_time_out
        old_pm_in = record.pm_time_in
        old_pm_out = record.pm_time_out

        # Update if a valid new objects exists, else if explicitly empty string from form update to None, 
        # but keep existing if not passed.
        if "am_in" in data:
            record.am_time_in = timezone.make_aware(datetime.combine(date_obj, am_in_obj)) if am_in_obj else None
        if "am_out" in data:
            record.am_time_out = timezone.make_aware(datetime.combine(date_obj, am_out_obj)) if am_out_obj else None
        if "pm_in" in data:
            record.pm_time_in = timezone.make_aware(datetime.combine(date_obj, pm_in_obj)) if pm_in_obj else None
        if "pm_out" in data:
            record.pm_time_out = timezone.make_aware(datetime.combine(date_obj, pm_out_obj)) if pm_out_obj else None
        record.save()

        return JsonResponse({"message": "Record updated successfully"})
    except Exception as e:
        return JsonResponse({"error": f"Error updating record: {str(e)}"}, status=400)


@csrf_exempt
def delete_record(request):
    if request.method != "POST":
        return JsonResponse({"error": "POST required"}, status=405)

    try:
        data = json.loads(request.body)
    except:
        return JsonResponse({"error": "Invalid JSON"}, status=400)

    student_id = data.get("student_id")
    record_id = data.get("record_id")

    if not all([student_id, record_id]):
        return JsonResponse({"error": "Missing required fields"}, status=400)

    try:
        record = Attendance.objects.get(id=record_id, student_id=student_id)
        record.delete()
        return JsonResponse({"message": "Record deleted successfully"})
    except Attendance.DoesNotExist:
        return JsonResponse({"error": "Record not found"}, status=404)


def download_dtr(request):
    import docx
    from datetime import datetime
    import calendar
    import os
    from django.conf import settings

    student_id = request.GET.get("student_id")
    if not student_id:
        return JsonResponse({"error": "Missing student_id"}, status=400)

    day_type = request.GET.get("day_type", "Regular")
    supervisor = request.GET.get("supervisor", "").strip()

    try:
        user = Intern.objects.get(student_id=student_id)
    except Intern.DoesNotExist:
        return JsonResponse({"error": "Intern not found"}, status=404)

    # Use current month and year
    now = timezone.localtime()
    year = now.year
    month = now.month
    month_name = calendar.month_name[month]
    month_str = f"{month_name} {year}"

    # Load Template
    template_path = os.path.join(settings.BASE_DIR, 'template', 'DTR_Template.docx')
    if not os.path.exists(template_path):
        return JsonResponse({"error": "Template not found"}, status=404)

    doc = docx.Document(template_path)

    # Helper to replace text while keeping underline and fixing font to prevent 2nd page spill
    def replace_paragraph(index, search, replacement, length=40, font_size=10, bold_name=True):
        if len(doc.paragraphs) > index:
            p = doc.paragraphs[index]
            p.clear()  # Clear existing runs
            
            from docx.shared import Pt
            
            if search:
                r1 = p.add_run(search + " ")
                r1.font.name = 'Times New Roman'
                r1.font.size = Pt(font_size)
            
            # Pad with non-breaking spaces to preserve underline width
            padded = replacement.upper().center(length, "\u00A0") if replacement else ("\u00A0" * length)
            run2 = p.add_run(padded)
            run2.underline = True
            run2.bold = bold_name
            run2.font.name = 'Times New Roman'
            run2.font.size = Pt(font_size)

    # Helper to insert tiny checks inline with native underline
    def check_inline(index, search, replacement):
        if len(doc.paragraphs) > index:
            p = doc.paragraphs[index]
            for run in p.runs:
                if search in run.text:
                    parts = run.text.split(search)
                    run.text = parts[0]
                    
                    from docx.shared import Pt
                    padded = replacement.center(len(search) + 2, "\u00A0")
                    r_check = p.add_run(padded)
                    r_check.underline = True
                    r_check.bold = True
                    r_check.font.name = 'Times New Roman'
                    r_check.font.size = Pt(11)
                    
                    if len(parts) > 1 and parts[1]:
                        p.add_run(parts[1])
                    break

    # Replace NAME
    replace_paragraph(6, "NAME:", user.name, 40, 10)
    replace_paragraph(39, "NAME:", user.name, 40, 10)
    
    # Replace MONTH
    replace_paragraph(8, "For the month of", month_str, 35, 10)
    replace_paragraph(41, "For the month of", month_str, 35, 10)
    
    # Replace SUPERVISOR (paragraphs 25 and 58 might be the underline lines, let's verify via lengths)
    if supervisor:
        replace_paragraph(25, "", supervisor, 35, 8, bold_name=True)
        replace_paragraph(58, "", supervisor, 35, 8, bold_name=True)
        
    # Checkmarks
    if day_type == "Regular":
        check_inline(10, "______", "✓")
        check_inline(43, "______", "✓")
    elif day_type == "Saturdays":
        check_inline(11, "_________", "✓")
        check_inline(44, "_________", "✓")

    records = Attendance.objects.filter(student_id=student_id, date__year=year, date__month=month).order_by('date')
    
    # Group records by day
    shifts_by_day = {}
    total_hours_month = 0
    for r in records:
        day = r.date.day
        if day not in shifts_by_day:
            shifts_by_day[day] = []
        shifts_by_day[day].append(r)

    # Populate Tables
    for t_idx in [0, 1]:  # Two identical tables
        if len(doc.tables) <= t_idx: break
        table = doc.tables[t_idx]
        
        for day in range(1, 32):
            if day > len(table.rows) - 2: break
            row = table.rows[day + 1] # Row 0 & 1 are headers
            
            if day in shifts_by_day:
                day_records = shifts_by_day[day]
                # Default empty
                am_in, am_out, pm_in, pm_out = "", "", "", ""
                
                rec1 = day_records[0]
                if rec1.am_time_in:
                    am_in = timezone.localtime(rec1.am_time_in).strftime("%I:%M")
                if rec1.am_time_out:
                    am_out = timezone.localtime(rec1.am_time_out).strftime("%I:%M")
                if rec1.pm_time_in:
                    pm_in = timezone.localtime(rec1.pm_time_in).strftime("%I:%M")
                if rec1.pm_time_out:
                    pm_out = timezone.localtime(rec1.pm_time_out).strftime("%I:%M")

                row.cells[1].text = am_in
                row.cells[2].text = am_out
                row.cells[3].text = pm_in
                row.cells[4].text = pm_out

    # Save to memory
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    filename = f"DTR_{user.name.replace(' ', '_')}_{month_name}_{year}.docx"
    response = FileResponse(buffer, as_attachment=True, filename=filename)
    return response


@csrf_exempt
def forgot_password(request):
    if request.method != "POST":
        return JsonResponse({"error": "POST request required"}, status=405)

    try:
        data = json.loads(request.body)
    except:
        return JsonResponse({"error": "Invalid JSON"}, status=400)

    student_id = data.get("student_id")
    email = data.get("email")
    new_password = data.get("new_password")

    if not all([student_id, email, new_password]):
        return JsonResponse({"error": "All fields are required"}, status=400)

    if len(new_password) < 8:
        return JsonResponse({"error": "Password must be at least 8 characters"}, status=400)

    try:
        user = Intern.objects.get(student_id=student_id, email=email)
        user.password = make_password(new_password)
        user.save()
        return JsonResponse({"message": "Password reset successfully"})
    except Intern.DoesNotExist:
        return JsonResponse({"error": "No account found with that Student ID and Email"}, status=404)


@csrf_exempt
def change_password(request):
    if request.method != "POST":
        return JsonResponse({"error": "POST request required"}, status=405)

    try:
        data = json.loads(request.body)
    except:
        return JsonResponse({"error": "Invalid JSON"}, status=400)

    student_id = data.get("student_id")
    current_password = data.get("current_password")
    new_password = data.get("new_password")

    if not all([student_id, current_password, new_password]):
        return JsonResponse({"error": "All fields are required"}, status=400)

    if len(new_password) < 8:
        return JsonResponse({"error": "New password must be at least 8 characters"}, status=400)

    try:
        user = Intern.objects.get(student_id=student_id)
        if not check_password(current_password, user.password):
            return JsonResponse({"error": "Current password is incorrect"}, status=401)

        user.password = make_password(new_password)
        user.save()
        return JsonResponse({"message": "Password changed successfully"})
    except Intern.DoesNotExist:
        return JsonResponse({"error": "Account not found"}, status=404)


@csrf_exempt
def update_profile(request):
    if request.method != "POST":
        return JsonResponse({"error": "POST request required"}, status=405)

    try:
        data = json.loads(request.body)
    except:
        return JsonResponse({"error": "Invalid JSON"}, status=400)

    student_id = data.get("student_id")
    name = data.get("name")
    email = data.get("email")

    if not student_id:
        return JsonResponse({"error": "Missing student_id"}, status=400)

    try:
        user = Intern.objects.get(student_id=student_id)

        if name:
            user.name = name
        if email:
            # Check if email is already taken by another user
            existing = Intern.objects.filter(email=email).exclude(student_id=student_id).first()
            if existing:
                return JsonResponse({"error": "Email is already in use"}, status=400)
            user.email = email

        user.save()
        return JsonResponse({"message": "Profile updated successfully", "name": user.name, "email": user.email})
    except Intern.DoesNotExist:
        return JsonResponse({"error": "Account not found"}, status=404)


def get_profile(request):
    student_id = request.GET.get("student_id")
    if not student_id:
        return JsonResponse({"error": "Missing student_id"}, status=400)

    try:
        user = Intern.objects.get(student_id=student_id)
        return JsonResponse({
            "name": user.name,
            "email": user.email,
            "student_id": user.student_id
        })
    except Intern.DoesNotExist:
        return JsonResponse({"error": "Account not found"}, status=404)


@csrf_exempt
def submit_report(request):
    if request.method != "POST":
        return JsonResponse({"error": "POST request required"}, status=405)

    student_id = request.POST.get("student_id")
    notes = request.POST.get("notes")
    images = request.FILES.getlist("images")

    if not student_id or not notes:
        return JsonResponse({"error": "Student ID and notes are required"}, status=400)

    try:
        intern = Intern.objects.get(student_id=student_id)
        
        # Create Report
        report = AccomplishmentReport.objects.create(student_id=student_id, notes=notes)
        
        # Save Images
        for image in images:
            AccomplishmentImage.objects.create(report=report, image=image)
        
        return JsonResponse({"message": "Report submitted successfully!", "report_id": report.id})
    except Intern.DoesNotExist:
        return JsonResponse({"error": "Account not found"}, status=404)
    except Exception as e:
        return JsonResponse({"error": str(e)}, status=500)


@csrf_exempt
def edit_report(request):
    if request.method != "POST":
        return JsonResponse({"error": "POST request required"}, status=405)

    try:
        data = json.loads(request.body)
    except Exception:
        return JsonResponse({"error": "Invalid JSON"}, status=400)

    student_id = data.get("student_id")
    report_id = data.get("report_id")
    notes = (data.get("notes") or "").strip()

    if not student_id or not report_id:
        return JsonResponse({"error": "student_id and report_id are required"}, status=400)
    if not notes:
        return JsonResponse({"error": "Notes cannot be empty"}, status=400)

    try:
        report = AccomplishmentReport.objects.get(id=report_id, student_id=student_id)
    except AccomplishmentReport.DoesNotExist:
        return JsonResponse({"error": "Report not found"}, status=404)

    report.notes = notes
    report.save(update_fields=["notes"])

    return JsonResponse({
        "message": "Report updated successfully",
        "report": {
            "id": report.id,
            "notes": report.notes,
        },
    })


@csrf_exempt
def delete_report(request):
    if request.method != "POST":
        return JsonResponse({"error": "POST request required"}, status=405)

    try:
        data = json.loads(request.body)
    except Exception:
        return JsonResponse({"error": "Invalid JSON"}, status=400)

    student_id = data.get("student_id")
    report_id = data.get("report_id")

    if not student_id or not report_id:
        return JsonResponse({"error": "student_id and report_id are required"}, status=400)

    try:
        report = AccomplishmentReport.objects.get(id=report_id, student_id=student_id)
    except AccomplishmentReport.DoesNotExist:
        return JsonResponse({"error": "Report not found"}, status=404)

    for image in report.images.all():
        if image.image:
            image.image.delete(save=False)

    report.delete()
    return JsonResponse({"message": "Report deleted successfully"})


def get_reports(request):
    student_id = request.GET.get("student_id")
    if not student_id:
        return JsonResponse({"error": "Missing student_id"}, status=400)

    try:
        from django.conf import settings
        reports = AccomplishmentReport.objects.filter(student_id=student_id).order_by("-created_at")
        
        results = []
        for r in reports:
            image_urls = [request.build_absolute_uri(settings.MEDIA_URL + str(img.image)) for img in r.images.all()]
            results.append({
                "id": r.id,
                "date": r.date.strftime("%b %d, %Y"),
                "time": timezone.localtime(r.created_at).strftime("%I:%M %p"),
                "notes": r.notes,
                "images": len(image_urls),
                "image_urls": image_urls
            })
            
        return JsonResponse({"reports": results})
    except Exception as e:
        return JsonResponse({"error": str(e)}, status=500)

